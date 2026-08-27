import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    """Establece el color de fondo de una celda en una tabla de Word."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece los márgenes internos (padding) de una celda."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def create_styled_table(doc, headers, data, col_widths=None):
    """Crea una tabla profesional con encabezado azul oscuro y filas alternadas."""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Estilo de Encabezado
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")  # Azul Corporativo LatinPyme
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "Calibri"
            run.font.size = Pt(10)

    # Filas de Datos
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F0F4F8" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)

    # Anchos de columna si se especifican
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    # Espaciado después de la tabla
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(12)
    return table

def build_word_document(output_path):
    doc = Document()

    # Configuración de Márgenes de Página
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── ESTILOS ──
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(45, 55, 72)

    # ── TÍTULO PRINCIPAL ──
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("ESCUELA LATINPYME")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(27, 54, 93)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(20)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle_p.add_run("Diccionario Técnico y Documentación Detallada de Base de Datos\nTablas, Llaves, Valores, Parámetros y Configuraciones")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(74, 85, 104)

    # Línea divisoria
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(16)
    p_div_run = p_div.add_run("―" * 55)
    p_div_run.font.color.rgb = RGBColor(200, 200, 200)

    # ── RESUMEN ARQUITECTÓNICO ──
    h1 = doc.add_heading("1. Resumen de Arquitectura de la Base de Datos", level=1)
    h1.runs[0].font.color.rgb = RGBColor(27, 54, 93)

    doc.add_paragraph(
        "El sistema de la Escuela LatinPyme utiliza una arquitectura híbrida de base de datos relacional y documental "
        "diseñada para alta disponibilidad, rendimiento en mensajería WhatsApp en tiempo real y trazabilidad de auditoría. "
        "Los componentes principales de la base de datos se desglosan a continuación:"
    )

    doc.add_paragraph(
        "• Motor Principal: PostgreSQL alojado en Supabase, accedido mediante el cliente oficial REST API (supabase-py) con la clave de servicio (service_role).\n"
        "• Normalización de Datos: Separación en tablas relacionales de Usuarios, Conversaciones, Mensajes (Normalizados con Idempotencia), Prompts de IA por Aula y Registros de Actividad.\n"
        "• Persistencia Local y Fallback: Archivo data/conversations.json para funcionamiento en modo offline o simulación, y config/config.json para parámetros del sistema e Inteligencia Artificial.\n"
        "• Mecanismo de Seguridad: Contraseñas cifradas con algoritmo Bcrypt y control de borrado lógico (Soft Delete) en todas las tablas mediante la columna deleted_at."
    )

    # ── DETALLE TABLA POR TABLA ──
    h1_2 = doc.add_heading("2. Detalle Meticuloso por Tablas y Campos", level=1)
    h1_2.runs[0].font.color.rgb = RGBColor(27, 54, 93)

    # ── TABLA 1: USERS ──
    h2_users = doc.add_heading("2.1. Tabla: users (Usuarios del Sistema)", level=2)
    h2_users.runs[0].font.color.rgb = RGBColor(43, 108, 176)
    doc.add_paragraph(
        "Almacena las cuentas de acceso al panel administrativo de la plataforma web. Soporta autenticación de asesores comerciales "
        "y administradores con asignación de roles y control de estado activo/inactivo."
    )

    headers_users = ["Campo / Llave", "Tipo de Dato", "Restricciones / Default", "Propósito y Funcionamiento"]
    widths_users = [1.3, 1.1, 1.5, 2.6]
    data_users = [
        ["id", "BIGSERIAL", "PRIMARY KEY, Auto", "Identificador único numérico e incremental asignado a cada usuario."],
        ["username", "TEXT", "UNIQUE, NOT NULL", "Nombre de usuario para iniciar sesión en el panel administrativo."],
        ["email", "TEXT", "UNIQUE, NOT NULL", "Correo electrónico institucional del usuario para notificaciones y seguridad."],
        ["password", "TEXT", "NOT NULL", "Hash de la contraseña generado con Bcrypt. Nunca almacena texto plano."],
        ["full_name", "TEXT", "NOT NULL", "Nombre completo del asesor o administrador visualizado en el panel y asignaciones."],
        ["role", "TEXT", "DEFAULT 'asesor'\nCHECK (admin/asesor)", "Rol de permisos: 'admin' (Administración completa) o 'asesor' (Gestión de chats y ventas)."],
        ["is_active", "BOOLEAN", "DEFAULT TRUE", "Estado de la cuenta. TRUE = Permite login. FALSE = Cuenta suspendida (desasigna chats activos)."],
        ["created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha y hora exacta en que se registró la cuenta de usuario."],
        ["last_login", "TIMESTAMPTZ", "NULL", "Timestamp del último acceso exitoso al sistema web."],
        ["deleted_at", "TIMESTAMPTZ", "DEFAULT NULL", "Marca de borrado lógico (Soft Delete). NULL = Usuario activo. Timestamp = Fecha de eliminación."]
    ]
    create_styled_table(doc, headers_users, data_users, widths_users)

    # ── TABLA 2: CONVERSATIONS ──
    h2_conv = doc.add_heading("2.2. Tabla: conversations (Conversaciones WhatsApp)", level=2)
    h2_conv.runs[0].font.color.rgb = RGBColor(43, 108, 176)
    doc.add_paragraph(
        "Registra cada cliente o participante que interactúa a través de WhatsApp. Almacena el estado de atención, "
        "metadatos de contacto, el asesor asignado y la bandera de requerimiento de atención humana."
    )

    headers_conv = ["Campo / Llave", "Tipo de Dato", "Restricciones / Default", "Propósito y Funcionamiento"]
    widths_conv = [1.3, 1.1, 1.5, 2.6]
    data_conv = [
        ["phone", "TEXT", "PRIMARY KEY", "Número telefónico del contacto en formato E.164 (+57...) o ID de Meta. Llave primaria."],
        ["assigned_to", "BIGINT", "FK users(id)\nON DELETE SET NULL", "ID del asesor asignado para atender la conversación. NULL si no está asignado."],
        ["data", "JSONB", "DEFAULT '{}', NOT NULL", "Objeto JSONB estructurado con metadatos extendidos del contacto y del aula virtual."],
        ["name", "TEXT", "NULL", "Nombre del cliente o contacto obtenido desde WhatsApp o la API de Aulas."],
        ["avatar", "TEXT", "NULL", "URL de la imagen de perfil o iniciales (ej: 'AM') para mostrar en la interfaz de chat."],
        ["unread", "INT", "DEFAULT 0", "Contador de mensajes no leídos pendientes por revisión del asesor."],
        ["last_message", "TEXT", "NULL", "Extracto o texto del último mensaje intercambiado en la conversación."],
        ["last_ts", "TIMESTAMPTZ", "NULL", "Fecha y hora de la última interacción recibida o enviada."],
        ["human_required", "BOOLEAN", "DEFAULT FALSE", "Bandera de escalamiento. TRUE si el cliente solicitó asesor humano o la IA lo derivó."],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Timestamp del último cambio realizado en el registro de la conversación."],
        ["deleted_at", "TIMESTAMPTZ", "DEFAULT NULL", "Marca de borrado lógico (Soft Delete) de la conversación."]
    ]
    create_styled_table(doc, headers_conv, data_conv, widths_conv)

    # Subsección para JSONB Data
    h3_jsonb = doc.add_heading("Estructura Interna del Campo JSONB 'data' en conversations:", level=3)
    h3_jsonb.runs[0].font.color.rgb = RGBColor(45, 55, 72)

    headers_jsonb = ["Llave JSONB", "Tipo", "Descripción y Ejemplo de Uso"]
    widths_jsonb = [1.8, 1.0, 3.7]
    data_jsonb = [
        ["last_trigger", "String", "Nombre del último disparador/flujo activado por el cliente (ej: 'SoporteOdoo', 'soluciones_patrocinador')."],
        ["wa_id", "String", "Identificador único de WhatsApp Cloud API asignado a la cuenta comercial."],
        ["real_phone", "String", "Número de teléfono real del cliente sin enmascaramiento de privacidad."],
        ["display_phone", "String", "Número formateado para presentación visual en el dashboard."],
        ["aula_info", "Object", "Objeto con información del aula del usuario: 'aula_id', 'aula_nombre', 'dominio', 'empresa_patrocinadora'."]
    ]
    create_styled_table(doc, headers_jsonb, data_jsonb, widths_jsonb)

    # ── TABLA 3: MESSAGES ──
    h2_msg = doc.add_heading("2.3. Tabla: messages (Historial de Mensajes Normalizado)", level=2)
    h2_msg.runs[0].font.color.rgb = RGBColor(43, 108, 176)
    doc.add_paragraph(
        "Almacena cada mensaje enviado o recibido en el sistema de manera individual y normalizada. Incorpora "
        "la llave wa_message_id para garantizar la idempotencia y evitar registros duplicados ante reintentos de Meta."
    )

    headers_msg = ["Campo / Llave", "Tipo de Dato", "Restricciones / Default", "Propósito y Funcionamiento"]
    widths_msg = [1.3, 1.1, 1.5, 2.6]
    data_msg = [
        ["id", "BIGSERIAL", "PRIMARY KEY, Auto", "ID único secuencial numérico de cada mensaje."],
        ["conversation_phone", "TEXT", "FK conversations(phone)\nON DELETE CASCADE", "Número de teléfono al cual pertenece este mensaje. Si la conversación se borra, sus mensajes se eliminan."],
        ["wa_message_id", "TEXT", "UNIQUE", "ID único de mensaje de Meta WhatsApp (wamid...). Garantiza Idempotencia de webhooks."],
        ["direction", "TEXT", "NOT NULL\nCHECK (in/out)", "Dirección del mensaje: 'in' = Entrante (Cliente), 'out' = Saliente (Asesor o IA)."],
        ["body", "TEXT", "DEFAULT '', NOT NULL", "Contenido del texto enviado o recibido en el mensaje."],
        ["ts", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha y hora exacta de transmisión o recepción del mensaje."],
        ["ia_generated", "BOOLEAN", "DEFAULT FALSE", "TRUE si el mensaje fue generado de forma automática por la IA Isabella; FALSE si fue redactado por humano."],
        ["disparador", "TEXT", "DEFAULT NULL", "Nombre del disparador de sistema ejecutado como resultado de este mensaje (ej: 'agente_taller_declaracion')."],
        ["wa_sent", "BOOLEAN", "DEFAULT TRUE", "Estado de entrega hacia la API de WhatsApp. TRUE = Enviado correctamente."],
        ["simulado", "BOOLEAN", "DEFAULT FALSE", "TRUE si el mensaje se originó en el entorno de pruebas/simulador interno."],
        ["created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha y hora de inserción en la base de datos."],
        ["deleted_at", "TIMESTAMPTZ", "DEFAULT NULL", "Marca de borrado lógico (Soft Delete) del mensaje."]
    ]
    create_styled_table(doc, headers_msg, data_msg, widths_msg)

    # ── TABLA 4: ACTIVITY_LOG ──
    h2_log = doc.add_heading("2.4. Tabla: activity_log (Registro de Auditoría de Actividad)", level=2)
    h2_log.runs[0].font.color.rgb = RGBColor(43, 108, 176)
    doc.add_paragraph(
        "Guarda la traza de auditoría de todas las operaciones realizadas por los asesores, administradores o procesos "
        "automatizados del sistema para control de calidad y seguridad."
    )

    headers_log = ["Campo / Llave", "Tipo de Dato", "Restricciones / Default", "Propósito y Funcionamiento"]
    widths_log = [1.3, 1.1, 1.5, 2.6]
    data_log = [
        ["id", "BIGSERIAL", "PRIMARY KEY, Auto", "ID único del registro de actividad."],
        ["user_id", "BIGINT", "FK users(id)\nON DELETE SET NULL", "ID del usuario que ejecutó la acción. NULL para procesos automáticos de sistema."],
        ["action", "TEXT", "NOT NULL", "Nombre clave del evento registrado (ej: 'LOGIN', 'ASSIGN_CONVERSATION', 'MIGRATION_PHASE_1')."],
        ["detail", "TEXT", "NULL", "Detalles adicionales en texto plano o JSON con el resultado o los datos del evento."],
        ["ts", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha y hora exacta en que se registró la actividad."]
    ]
    create_styled_table(doc, headers_log, data_log, widths_log)

    # ── TABLA 5: AI_PROMPTS ──
    h2_prompts = doc.add_heading("2.5. Tabla: ai_prompts (Prompts de IA por Aula Virtual)", level=2)
    h2_prompts.runs[0].font.color.rgb = RGBColor(43, 108, 176)
    doc.add_paragraph(
        "Permite personalizar los comportamientos, reglas de negocio y tono de respuesta de la IA Isabella "
        "de acuerdo a la empresa o aula virtual específica en la que se encuentra el usuario."
    )

    headers_prompts = ["Campo / Llave", "Tipo de Dato", "Restricciones / Default", "Propósito y Funcionamiento"]
    widths_prompts = [1.3, 1.1, 1.5, 2.6]
    data_prompts = [
        ["id", "BIGSERIAL", "PRIMARY KEY, Auto", "ID único del registro de prompt por aula."],
        ["aula_id", "TEXT", "NOT NULL, UNIQUE", "Identificador único del aula virtual (ej: '2', '3', '115')."],
        ["aula_nombre", "TEXT", "NOT NULL", "Nombre descriptivo de la empresa o aula del cliente (ej: 'AULA EQUIDAD ARL')."],
        ["prompt", "TEXT", "NOT NULL", "Directivas y reglas del sistema que la IA ejecutará para los usuarios de esta aula."],
        ["activo", "BOOLEAN", "DEFAULT TRUE", "TRUE = Prompt del aula habilitado. FALSE = Usar prompt global por defecto."],
        ["created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha de creación del prompt en el sistema."],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha de la última modificación realizada al prompt."],
        ["deleted_at", "TIMESTAMPTZ", "DEFAULT NULL", "Marca de borrado lógico (Soft Delete) del prompt."]
    ]
    create_styled_table(doc, headers_prompts, data_prompts, widths_prompts)

    # ── TABLA 6: AI_CONVERSATION_CONFIG ──
    h2_aiconf = doc.add_heading("2.6. Tabla: ai_conversation_config (Estado de IA por Conversación)", level=2)
    h2_aiconf.runs[0].font.color.rgb = RGBColor(43, 108, 176)
    doc.add_paragraph(
        "Permite pausar o reactivar el bot de Inteligencia Artificial para un chat específico de forma individual, "
        "facilitando que un asesor humano tome el control total del canal sin interferencias de la IA."
    )

    headers_aiconf = ["Campo / Llave", "Tipo de Dato", "Restricciones / Default", "Propósito y Funcionamiento"]
    widths_aiconf = [1.3, 1.1, 1.5, 2.6]
    data_aiconf = [
        ["phone", "TEXT", "PRIMARY KEY", "Número telefónico del contacto en formato E.164. Llave primaria."],
        ["ai_enabled", "BOOLEAN", "DEFAULT TRUE", "TRUE = La IA Isabella responde automáticamente. FALSE = La IA está en silencio (atención humana)."],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha y hora del último cambio en el estado de la IA para este chat."],
        ["deleted_at", "TIMESTAMPTZ", "DEFAULT NULL", "Marca de borrado lógico del registro."]
    ]
    create_styled_table(doc, headers_aiconf, data_aiconf, widths_aiconf)

    # ── SECCIÓN 3: ÍNDICES ──
    h1_3 = doc.add_heading("3. Índices de Rendimiento y Optimización de la BD", level=1)
    h1_3.runs[0].font.color.rgb = RGBColor(27, 54, 93)
    doc.add_paragraph(
        "Para garantizar alta velocidad de respuesta y escalabilidad ante miles de conversaciones simultáneas, "
        "la base de datos utiliza índices parciales estratégicos que filtran registros activos (WHERE deleted_at IS NULL):"
    )

    headers_idx = ["Nombre del Índice", "Tabla", "Columnas / Condición", "Beneficio de Rendimiento"]
    widths_idx = [1.8, 1.2, 1.8, 1.7]
    data_idx = [
        ["idx_conversations_deleted_at", "conversations", "deleted_at (WHERE IS NULL)", "Búsqueda relámpago de conversaciones activas."],
        ["idx_conversations_assigned_to", "conversations", "assigned_to (WHERE IS NULL)", "Filtrado eficiente de chats asignados a cada asesor."],
        ["idx_messages_conv_ts", "messages", "conversation_phone, ts DESC (WHERE IS NULL)", "Carga instantánea de los últimos mensajes por chat."],
        ["idx_messages_deleted_at", "messages", "deleted_at (WHERE IS NULL)", "Exclusión de mensajes eliminados en consultas."],
        ["idx_users_deleted_at", "users", "deleted_at (WHERE IS NULL)", "Listado optimizado de asesores activos."],
        ["idx_users_username", "users", "username (WHERE IS NULL)", "Autenticación de login de alta velocidad."],
        ["idx_activity_log_user_ts", "activity_log", "user_id, ts DESC", "Generación rápida de reportes de auditoría por usuario."],
        ["idx_ai_prompts_aula_id", "ai_prompts", "aula_id (WHERE IS NULL)", "Obtención inmediata del prompt por aula durante el flujo de IA."]
    ]
    create_styled_table(doc, headers_idx, data_idx, widths_idx)

    # ── SECCIÓN 4: PARÁMETROS Y CONFIGURACIÓN JSON ──
    h1_4 = doc.add_heading("4. Parámetros de Archivo de Configuración (config/config.json y .env)", level=1)
    h1_4.runs[0].font.color.rgb = RGBColor(27, 54, 93)
    doc.add_paragraph(
        "El sistema combina variables de entorno (.env) con el archivo JSON config/config.json para administrar "
        "las credenciales de servicios externos, parámetros del agente de IA y llaves de seguridad:"
    )

    headers_cfg = ["Parámetro / Llave", "Tipo", "Valor por Defecto / Ejemplo", "Descripción y Uso en el Sistema"]
    widths_cfg = [1.8, 0.9, 1.7, 2.1]
    data_cfg = [
        ["OPENAI_API_KEY", "String", "sk-proj-...", "Clave API de OpenAI utilizada para generar respuestas inteligentes con la IA Isabella."],
        ["META_WA_TOKEN", "String", "tu_token_aqui", "Token Bearer de autenticación para enviar mensajes desde WhatsApp Cloud API."],
        ["META_WA_PHONE_NUMBER_ID", "String", "tu_phone_number_id", "Identificador único del número telefónico empresarial en Meta Business."],
        ["META_WA_BUSINESS_ACCOUNT_ID", "String", "tu_business_account_id", "ID de la cuenta comercial de la empresa en el Meta Business Manager."],
        ["META_WA_WEBHOOK_VERIFY_TOKEN", "String", "latinpyme_secret_token", "Token de validación utilizado durante la suscripción del Webhook de WhatsApp."],
        ["META_WA_TEMPLATE_NAME", "String", "plantilla_curso_latinpyme", "Nombre de la plantilla de mensaje HSM aprobada para envíos masivos."],
        ["AULAS_API_URL", "String", "https://capacitacion.../getUserPhone", "Endpoint HTTP de la API externa para obtener datos del alumno por teléfono."],
        ["AULAS_API_TOKEN", "String", "topchat-token-access-2024", "Token Bearer para autenticar las peticiones a la API de Aulas."],
        ["AI_ASSISTANT_ENABLED", "Boolean", "true", "Control global del bot de IA: 'true' (IA activa), 'false' (IA totalmente apagada)."],
        ["AI_AUTO_REPLY", "Boolean", "true", "Permite que la IA envíe respuestas automáticas sin intervención del asesor."],
        ["AI_MASTER_PROMPT", "String", "Texto extenso de instrucciones", "Prompt maestro de Isabella con directivas B2B, horarios, promociones y disparadores."],
        ["AI_DEFAULT_PROMPT", "String", "Texto por defecto", "Prompt básico de contingencia si no hay prompt configurado para un aula."],
        ["OPENAI_MODEL", "String", "gpt-4o-mini", "Modelo de lenguaje configurado para procesar las respuestas de la IA."],
        ["SUPABASE_URL", "String", "https://xyz.supabase.co", "URL del endpoint HTTPS de la instancia de Supabase PostgreSQL."],
        ["SUPABASE_SECRET_KEY", "String", "eyJhbGci...", "Clave de servicio (service_role) con acceso total a la base de datos (Bypasses RLS)."]
    ]
    create_styled_table(doc, headers_cfg, data_cfg, widths_cfg)

    # ── PIE Y FINAL ──
    doc.add_paragraph().paragraph_format.space_before = Pt(16)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_end = p_end.add_run("― Fin del Documento Técnico — Escuela LatinPyme ―")
    run_end.font.italic = True
    run_end.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(output_path)
    print(f"[OK] Documento Word generado exitosamente en: {output_path}")

if __name__ == "__main__":
    out_file = os.path.join(r"c:\Users\LATINPYMES\OneDrive\Desktop\proyecto_latinpyme", "Documentacion_Base_de_Datos_LatinPyme.docx")
    build_word_document(out_file)
